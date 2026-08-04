from pathlib import Path
from math import atan2, cos, sin, pi
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
ARCH_PNG = OUT_DIR / "AV_GRU_Autoencoder_Architecture.png"
FLOW_PNG = OUT_DIR / "AV_GRU_Autoencoder_End_to_End_Flow.png"
DOCX_PATH = OUT_DIR / "AV_GRU_Autoencoder_Model_Flow.docx"

INK = (21, 35, 48)
MUTED = (83, 96, 111)
BLUE = (46, 116, 181)
BLUE_DARK = (31, 77, 120)
BLUE_FILL = (232, 238, 245)
GREEN_FILL = (228, 242, 237)
GOLD_FILL = (250, 241, 217)
GRAY_FILL = (244, 246, 248)
BORDER = (128, 145, 163)
WHITE = (255, 255, 255)


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


TITLE_FONT = font(34, bold=True)
BOX_FONT = font(25, bold=True)
BODY_FONT = font(21)
SMALL_FONT = font(18)


def text_size(draw, text, fnt):
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw, text, fnt, max_width):
    lines = []
    for raw_line in text.split("\n"):
        words = raw_line.split()
        if not words:
            lines.append("")
            continue
        line = words[0]
        for word in words[1:]:
            trial = f"{line} {word}"
            if text_size(draw, trial, fnt)[0] <= max_width:
                line = trial
            else:
                lines.append(line)
                line = word
        lines.append(line)
    return lines


def rounded_box(draw, xy, title, details="", fill=GRAY_FILL, outline=BORDER):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=22, fill=fill, outline=outline, width=4)
    title_lines = wrap_text(draw, title, BOX_FONT, x2 - x1 - 42)
    detail_lines = wrap_text(draw, details, BODY_FONT, x2 - x1 - 42) if details else []
    line_gap = 8
    total_h = len(title_lines) * 31 + len(detail_lines) * 27 + (line_gap if detail_lines else 0)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in title_lines:
        w, _ = text_size(draw, line, BOX_FONT)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=BOX_FONT, fill=INK)
        y += 31
    if detail_lines:
        y += line_gap
    for line in detail_lines:
        w, _ = text_size(draw, line, BODY_FONT)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=BODY_FONT, fill=MUTED)
        y += 27


def arrow(draw, start, end, color=BLUE_DARK, width=6):
    draw.line([start, end], fill=color, width=width)
    angle = atan2(end[1] - start[1], end[0] - start[0])
    size = 20
    left = (
        end[0] - size * cos(angle - pi / 6),
        end[1] - size * sin(angle - pi / 6),
    )
    right = (
        end[0] - size * cos(angle + pi / 6),
        end[1] - size * sin(angle + pi / 6),
    )
    draw.polygon([end, left, right], fill=color)


def draw_caption(draw, text, xy):
    x, y = xy
    draw.text((x, y), text, font=SMALL_FONT, fill=MUTED)


def build_architecture_png():
    img = Image.new("RGB", (1800, 1240), WHITE)
    draw = ImageDraw.Draw(img)
    draw.text((70, 48), "AV-GRU Autoencoder Architecture", font=TITLE_FONT, fill=INK)
    draw.text(
        (70, 95),
        "Motion-only encoder with optional audio-guided training branch",
        font=BODY_FONT,
        fill=MUTED,
    )

    boxes = {
        "input": (70, 180, 390, 330),
        "encoder": (500, 180, 820, 330),
        "pool": (930, 180, 1250, 330),
        "latent": (1360, 180, 1680, 330),
        "repeat": (500, 520, 820, 660),
        "decoder": (500, 760, 820, 900),
        "motion_head": (500, 1000, 820, 1120),
        "recon": (930, 1000, 1250, 1120),
        "audio_head": (1360, 520, 1680, 660),
        "audio_pred": (1360, 760, 1680, 900),
        "loss": (930, 520, 1250, 660),
    }

    rounded_box(draw, boxes["input"], "Motion window", "B x T x motion_dim\ndefault B x 60 x 126", BLUE_FILL)
    rounded_box(draw, boxes["encoder"], "GRU encoder", "bidirectional when configured\nhidden_dim x directions", BLUE_FILL)
    rounded_box(draw, boxes["pool"], "Mean pool", "average encoded states\nover time", BLUE_FILL)
    rounded_box(draw, boxes["latent"], "Latent z", "LayerNorm + Linear\ndefault B x 24", GREEN_FILL)
    rounded_box(draw, boxes["repeat"], "Repeat z", "same latent vector\nat each time step", GREEN_FILL)
    rounded_box(draw, boxes["decoder"], "GRU decoder", "latent sequence ->\nhidden states", BLUE_FILL)
    rounded_box(draw, boxes["motion_head"], "Motion head", "Linear hidden_dim\n-> motion_dim", BLUE_FILL)
    rounded_box(draw, boxes["recon"], "Motion recon", "B x T x motion_dim", BLUE_FILL)
    rounded_box(draw, boxes["audio_head"], "Audio head", "Linear -> ReLU -> Linear", GOLD_FILL)
    rounded_box(draw, boxes["audio_pred"], "Audio pred", "B x audio_dim\ntraining target only", GOLD_FILL)
    rounded_box(draw, boxes["loss"], "Loss", "motion MSE + gated audio MSE\n+ latent smoothness", GRAY_FILL)

    arrow(draw, (390, 255), (500, 255))
    arrow(draw, (820, 255), (930, 255))
    arrow(draw, (1250, 255), (1360, 255))
    arrow(draw, (1520, 330), (1520, 520))
    arrow(draw, (1520, 660), (1520, 760))
    arrow(draw, (1360, 255), (820, 520))
    arrow(draw, (660, 660), (660, 760))
    arrow(draw, (660, 900), (660, 1000))
    arrow(draw, (820, 1060), (930, 1060))
    arrow(draw, (1250, 1060), (1090, 660), color=(97, 112, 128), width=4)
    arrow(draw, (1520, 900), (1250, 590), color=(97, 112, 128), width=4)
    arrow(draw, (1520, 330), (1250, 555), color=(97, 112, 128), width=4)

    draw_caption(draw, "Runtime uses the encoder path to produce z from live motion windows.", (70, 1145))
    draw_caption(draw, "The audio branch shapes z during training when audio loss weight and audio_quality are nonzero.", (70, 1178))
    img.save(ARCH_PNG, quality=95)


def build_flow_png():
    img = Image.new("RGB", (1800, 880), WHITE)
    draw = ImageDraw.Draw(img)
    draw.text((70, 48), "Stage 2 Training, Embedding, Clustering, and Runtime Export", font=TITLE_FONT, fill=INK)

    y = 210
    w = 235
    h = 145
    gap = 42
    labels = [
        ("Stage 1 CSVs", "hand labels +\nlandmark columns"),
        ("build_windows.py", "normalize, order hands,\nmake fixed windows"),
        (".npz windows", "motion, audio,\naudio_quality"),
        ("train.py", "fit scaler + GRU AE\nsave best checkpoint"),
        ("embed.py", "scale motion, encode,\nwrite latent z"),
        ("cluster.py", "UMAP view +\nHDBSCAN/KMeans"),
    ]
    x = 70
    coords = []
    for i, (title, details) in enumerate(labels):
        xy = (x + i * (w + gap), y, x + i * (w + gap) + w, y + h)
        coords.append(xy)
        fill = BLUE_FILL if i < 4 else GREEN_FILL
        rounded_box(draw, xy, title, details, fill)
        if i > 0:
            prev = coords[i - 1]
            arrow(draw, (prev[2], y + h / 2), (xy[0], y + h / 2))

    export_xy = (1195, 550, 1510, 710)
    runtime_xy = (70, 550, 430, 710)
    rounded_box(draw, export_xy, "export_for_runtime.py", "encoder.pt, scalers,\ncluster model, config", GOLD_FILL)
    rounded_box(draw, runtime_xy, "Stage 3 runtime", "same feature contract\nlive windows -> z -> cluster", GOLD_FILL)
    arrow(draw, ((coords[3][0] + coords[3][2]) / 2, coords[3][3]), (1350, 550))
    arrow(draw, ((coords[5][0] + coords[5][2]) / 2, coords[5][3]), (1350, 550))
    arrow(draw, (1195, 630), (430, 630))

    draw_caption(draw, "The scaler and feature contract are as important as the checkpoint: Stage 3 must reproduce training-time inputs.", (70, 800))
    img.save(FLOW_PNG, quality=95)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, (46, 116, 181), 18, 10),
        ("Heading 2", 13, (46, 116, 181), 14, 7),
        ("Heading 3", 12, (31, 77, 120), 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_para(doc, text="", style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_run_font(r, size=9, color=(83, 96, 111), italic=True)


def add_config_table(doc):
    data = [
        ["Config", "Sequence", "Motion dim", "Audio dim", "Hidden", "Latent", "Layers", "Audio loss"],
        ["default.yaml", "60", "126", "12", "128", "24", "1", "0.35"],
        ["full_av.yaml", "60", "126", "12", "128", "24", "2", "0.35"],
        ["motion_only.yaml", "60", "126", "1", "128", "24", "2", "0.0"],
        ["small_test.yaml", "60", "126", "1", "64", "16", "1", "0.0"],
        ["small_test_seq15.yaml", "15", "126", "1", "64", "16", "1", "0.0"],
    ]
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [
        Inches(1.72),
        Inches(0.58),
        Inches(0.68),
        Inches(0.64),
        Inches(0.58),
        Inches(0.58),
        Inches(0.50),
        Inches(0.72),
    ]
    set_table_width(table, widths)
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=8.5, bold=(r_idx == 0), color=(21, 35, 48))
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
            if c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def add_artifacts_table(doc):
    data = [
        ["Artifact", "Purpose"],
        ["best_model.pt / encoder.pt", "Model checkpoint containing the trained AVGRUAutoencoder weights and config."],
        ["feature_scaler.joblib", "StandardScaler fit on training motion frames; required before encoding."],
        ["embeddings.csv / embeddings.npy", "Per-window latent gesture vectors generated by model.encode(...)."],
        ["cluster_model.joblib", "HDBSCAN or KMeans model trained on standardized latent vectors."],
        ["embedding_scaler.joblib", "Scaler applied to latent vectors before clustering."],
        ["runtime_model_config.json", "Feature contract and model dimensions that Stage 3 must follow."],
    ]
    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"
    widths = [Inches(2.05), Inches(4.45)]
    set_table_width(table, widths)
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=9.5, bold=(r_idx == 0), color=(21, 35, 48))
            if r_idx == 0:
                set_cell_shading(cell, "E8EEF5")
    doc.add_paragraph()


def build_docx():
    doc = Document()
    configure_doc(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    title = p.add_run("AV-GRU Autoencoder Model and Flow")
    set_run_font(title, size=23, color=(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    subtitle = p.add_run("Stage 2 architecture, training path, embedding flow, and runtime handoff")
    set_run_font(subtitle, size=12, color=(83, 96, 111), italic=True)

    add_para(doc, "Source reviewed: /Users/spaceboxcarbonite/Downloads/av_GRU_autoencoder")

    doc.add_heading("Purpose", level=1)
    add_para(
        doc,
        "This Stage 2 package turns Stage 1 hand-landmark CSVs into fixed-length motion windows, trains a GRU autoencoder, converts each window into a compact latent gesture vector, clusters those latent vectors, and exports the trained artifacts for Stage 3 runtime use."
    )
    add_para(
        doc,
        "Key point: the encoder input is motion only. Audio is an optional auxiliary training target, gated by each window's audio_quality value. The deployed runtime can therefore operate from camera/MediaPipe hand landmarks without live audio features.",
        bold_lead="Key point:",
    )

    doc.add_heading("End-to-End Flow", level=1)
    doc.add_picture(str(FLOW_PNG), width=Inches(6.45))
    add_caption(doc, "Figure 1. Data and artifact flow from Stage 1 landmark CSVs to Stage 3 runtime package.")

    doc.add_heading("Model Architecture", level=1)
    doc.add_picture(str(ARCH_PNG), width=Inches(6.45))
    add_caption(doc, "Figure 2. Forward pass through the motion encoder, latent vector, motion decoder, and audio prediction branch.")

    doc.add_heading("Data Contract", level=1)
    add_para(
        doc,
        "The default motion vector is right hand 63 values plus left hand 63 values, for 126 total features. Each hand has 21 MediaPipe landmarks with 3 coordinates. The builder uses handedness labels when available so the model sees a stable right-then-left order instead of raw detection order."
    )
    add_bullets(doc, [
        "Missing hands are represented with zeros.",
        "Each hand is normalized relative to the wrist landmark.",
        "Scale normalization uses landmark 9 by default.",
        "Optional velocity features can be appended, but the default configs keep include_velocity false.",
        "Each .npz window stores motion, audio, audio_quality, source metadata, and a serialized feature contract.",
    ])

    doc.add_heading("Forward Pass", level=1)
    add_para(
        doc,
        "For a batch shaped B x T x motion_dim, the encoder GRU emits B x T x hidden_dim*directions. The model averages those encoded states across time and projects the result through LayerNorm plus Linear to produce the latent gesture vector z."
    )
    add_para(
        doc,
        "The decoder receives z repeated across T time steps, then reconstructs the full motion sequence through a decoder GRU and linear motion head. The audio head separately maps z through Linear, ReLU, and Linear layers to predict an audio summary vector."
    )
    add_para(
        doc,
        "Runtime note: Stage 3 only needs the encoder path and preprocessing artifacts. The audio prediction branch is useful during training when real audio guidance is enabled, but it is not required for live inference.",
        bold_lead="Runtime note:",
    )

    doc.add_heading("Training Objective", level=1)
    add_para(
        doc,
        "The total loss is a weighted sum of motion reconstruction MSE, gated audio prediction MSE, and a latent smoothness regularizer."
    )
    add_bullets(doc, [
        "Motion reconstruction loss compares motion_recon against the scaled input motion sequence.",
        "Audio prediction loss compares audio_pred against the audio target; if the target is time-varying, it is averaged over time first.",
        "The audio squared error is multiplied by audio_quality, so placeholder or poor audio can have little influence.",
        "Motion-only configs set audio_prediction_weight to 0.0, which disables the audio branch's contribution.",
        "The smoothness term compares adjacent latent vectors in the current batch. With shuffled training batches, it is a light regularizer rather than guaranteed source-video temporal smoothing.",
    ])

    doc.add_heading("Training Flow", level=1)
    add_para(
        doc,
        "Training discovers the manifest, validates that the first window matches the configured sequence length and feature dimensions, creates train/validation/test splits, fits a StandardScaler on flattened training motion frames, and applies that scaler to every batch before the model sees it."
    )
    add_para(
        doc,
        "The model is optimized with AdamW, gradient clipping, early stopping on validation total loss, and checkpointing of the best validation model."
    )

    doc.add_heading("Embedding, Clustering, and Evaluation", level=1)
    add_para(
        doc,
        "After training, embed.py reloads the feature scaler and best checkpoint, scales each motion window, calls model.encode(...), and writes the latent vectors to embeddings.npy and embeddings.csv. The clustering step standardizes the latent columns, fits UMAP for two-dimensional visualization, then clusters the standardized latent vectors with HDBSCAN or KMeans. HDBSCAN is the default."
    )
    add_para(
        doc,
        "The evaluation script summarizes total windows, non-noise clusters, noise windows, and clustering metrics such as silhouette and Davies-Bouldin when enough valid clusters exist."
    )

    doc.add_heading("Runtime Export", level=1)
    add_para(
        doc,
        "export_for_runtime.py packages the checkpoint, motion scaler, clustering artifacts, cluster names, and a compact runtime_model_config.json file. The runtime config records the feature contract and model dimensions that Stage 3 must reproduce."
    )
    add_artifacts_table(doc)

    doc.add_page_break()
    doc.add_heading("Configuration Profiles", level=1)
    add_config_table(doc)
    add_para(
        doc,
        "The code contains a use_audio_guidance config field, but the active loss behavior is controlled by audio_prediction_weight and audio_quality. In the current implementation, setting audio_prediction_weight to 0.0 is the effective way to disable audio guidance."
    )

    doc.add_heading("Implementation Notes", level=1)
    add_bullets(doc, [
        "from_latent is defined in the model but is not used by the current forward method.",
        "audio_quality_threshold appears in config files, but the shown loss code directly uses per-window audio_quality rather than thresholding it.",
        "Runtime correctness depends on keeping hand order, normalization, scale landmark, sequence length, feature dimension, and scaler identical between training and Stage 3.",
        "The architecture is practical for live camera-driven interaction because the encoder consumes motion only, even when training uses audio as additional supervision.",
    ])

    doc.add_heading("Source Files Reviewed", level=1)
    add_bullets(doc, [
        "models/av_gru_autoencoder.py",
        "src/losses.py",
        "src/dataset.py",
        "build_windows.py, train.py, embed.py, cluster.py, evaluate.py, export_for_runtime.py",
        "configs/default.yaml, full_av.yaml, motion_only.yaml, small_test.yaml, small_test_seq15.yaml",
    ])

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_architecture_png()
    build_flow_png()
    build_docx()
    print(DOCX_PATH)
